import io
from pathlib import Path

from docx import Document
from pypdf import PdfReader

from app.api.errors import ApiError
from app.config import Settings
from app.documents.models import ExtractedDocument


def extract_document(file_path: Path, file_type: str, settings: Settings) -> ExtractedDocument:
    """Dispatch extraction by normalized file type."""

    if file_type in {"txt", "md"}:
        return extract_text_document(file_path, settings)
    if file_type == "docx":
        return extract_docx_document(file_path, settings)
    if file_type == "pdf":
        return extract_pdf_document(file_path, settings)
    raise ApiError(415, "UNSUPPORTED_FILE_TYPE", "Bu fayl turi qo'llab-quvvatlanmaydi.")


def normalize_text(text: str) -> str:
    """Normalize newlines and excessive blank lines."""

    text = text.replace("\r\n", "\n").replace("\r", "\n").lstrip("\ufeff")
    lines = [line.rstrip() for line in text.split("\n")]
    normalized: list[str] = []
    blank_count = 0
    for line in lines:
        if line == "":
            blank_count += 1
            if blank_count > 1:
                continue
        else:
            blank_count = 0
        normalized.append(line)
    return "\n".join(normalized).strip()


def extract_text_document(file_path: Path, settings: Settings) -> ExtractedDocument:
    """Extract normalized UTF-8 text from TXT or MD files."""

    raw = file_path.read_bytes()
    if b"\x00" in raw:
        raise ApiError(422, "INVALID_TEXT_ENCODING", "Matn faylda binary ma'lumot topildi.")
    try:
        text = raw.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise ApiError(422, "INVALID_TEXT_ENCODING", "Matn fayl UTF-8 formatida emas.") from exc
    normalized = normalize_text(text)
    if len(normalized) > settings.max_extracted_chars:
        raise ApiError(422, "EXTRACTED_TEXT_TOO_LARGE", "Extract qilingan matn juda katta.")
    return ExtractedDocument(
        text=normalized,
        char_count=len(normalized),
        page_count=None,
        status="ready" if normalized else "no_text",
        warning_code=None if normalized else "NO_EXTRACTABLE_TEXT",
    )


def extract_docx_document(file_path: Path, settings: Settings) -> ExtractedDocument:
    """Extract paragraphs and tables from DOCX files."""

    try:
        document = Document(file_path)
    except Exception as exc:
        raise ApiError(422, "INVALID_DOCX", "DOCX faylni ochib bo'lmadi.") from exc

    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    normalized = normalize_text("\n\n".join(parts))
    if len(normalized) > settings.max_extracted_chars:
        raise ApiError(422, "EXTRACTED_TEXT_TOO_LARGE", "Extract qilingan matn juda katta.")
    return ExtractedDocument(
        text=normalized,
        char_count=len(normalized),
        page_count=None,
        status="ready" if normalized else "no_text",
        warning_code=None if normalized else "NO_EXTRACTABLE_TEXT",
    )


def extract_pdf_document(file_path: Path, settings: Settings) -> ExtractedDocument:
    """Extract text from a PDF file with page and size limits."""

    try:
        reader = PdfReader(file_path, strict=False)
    except Exception as exc:
        raise ApiError(422, "INVALID_PDF", "PDF faylni ochib bo'lmadi.") from exc

    if reader.is_encrypted:
        try:
            opened = reader.decrypt("")
        except Exception:
            opened = 0
        if opened == 0:
            raise ApiError(422, "PDF_ENCRYPTED", "Parol bilan himoyalangan PDF qabul qilinmaydi.")

    if len(reader.pages) > settings.max_pdf_pages:
        raise ApiError(422, "PDF_PAGE_LIMIT_EXCEEDED", "PDF sahifalari soni limitdan oshdi.")

    page_texts: list[str] = []
    for page in reader.pages:
        contents = page.get_contents()
        if contents is not None:
            if isinstance(contents, list):
                raw_size = sum(len(item.get_data()) for item in contents if item is not None)
            else:
                raw_size = len(contents.get_data())
            if raw_size > settings.max_pdf_page_content_mb * 1024 * 1024:
                raise ApiError(422, "PDF_CONTENT_TOO_LARGE", "PDF sahifa content hajmi juda katta.")
        text = page.extract_text() or ""
        page_texts.append(text.strip())
        if sum(len(item) for item in page_texts) > settings.max_extracted_chars:
            raise ApiError(422, "EXTRACTED_TEXT_TOO_LARGE", "Extract qilingan matn juda katta.")

    normalized = normalize_text("\n\n".join(part for part in page_texts if part))
    return ExtractedDocument(
        text=normalized,
        char_count=len(normalized),
        page_count=len(reader.pages),
        status="ready" if normalized else "no_text",
        warning_code=None if normalized else "OCR_REQUIRED",
    )
