import zipfile
from pathlib import Path

import pytest
from docx import Document
from pypdf import PdfWriter

from app.api.errors import ApiError
from app.documents.extractor import extract_document
from app.documents.isolated_extraction import extract_document_isolated
from app.documents.validator import validate_docx_archive
from tests.conftest import build_settings


def create_docx(path: Path, paragraphs: list[str], table_rows: list[list[str]] | None = None) -> None:
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for row_index, row in enumerate(table_rows):
            for column_index, value in enumerate(row):
                table.cell(row_index, column_index).text = value
    document.save(path)


def create_pdf(path: Path, encrypted: bool = False, pages: int = 1) -> None:
    writer = PdfWriter()
    for _ in range(pages):
        writer.add_blank_page(width=200, height=200)
    if encrypted:
        writer.encrypt("secret")
    with open(path, "wb") as handle:
        writer.write(handle)


class InlineProcess:
    def __init__(self, *, target, args):
        self._target = target
        self._args = args
        self._alive = False
        self.pid = 4242
        self.terminated = False
        self.killed = False

    def start(self):
        self._alive = True
        self._target(*self._args)
        self._alive = False

    def is_alive(self):
        return self._alive

    def join(self, timeout=None):
        return None

    def terminate(self):
        self.terminated = True
        self._alive = False

    def kill(self):
        self.killed = True
        self._alive = False


class HangingProcess:
    def __init__(self, *, target, args):
        self.pid = 4343
        self._alive = False
        self.terminated = False
        self.killed = False

    def start(self):
        self._alive = True

    def is_alive(self):
        return self._alive

    def join(self, timeout=None):
        return None

    def terminate(self):
        self.terminated = True
        self._alive = False

    def kill(self):
        self.killed = True
        self._alive = False


def test_extract_txt_document(tmp_path: Path) -> None:
    settings = build_settings(tmp_path)
    path = tmp_path / "sample.txt"
    path.write_text("A\r\n\r\n\r\nB  \n", encoding="utf-8")
    extracted = extract_document(path, "txt", settings)
    assert extracted.text == "A\n\nB"
    assert extracted.status == "ready"


def test_extract_md_document(tmp_path: Path) -> None:
    settings = build_settings(tmp_path)
    path = tmp_path / "sample.md"
    path.write_text("# Title\n\nText", encoding="utf-8")
    extracted = extract_document(path, "md", settings)
    assert "Title" in extracted.text


def test_extract_docx_document(tmp_path: Path) -> None:
    settings = build_settings(tmp_path)
    path = tmp_path / "sample.docx"
    create_docx(path, ["Heading", "Paragraph"], [["A", "B"]])
    extracted = extract_document(path, "docx", settings)
    assert "Heading" in extracted.text
    assert "A | B" in extracted.text


def test_extract_blank_pdf_returns_no_text(tmp_path: Path) -> None:
    settings = build_settings(tmp_path)
    path = tmp_path / "blank.pdf"
    create_pdf(path)
    extracted = extract_document(path, "pdf", settings)
    assert extracted.status == "no_text"
    assert extracted.warning_code == "OCR_REQUIRED"


def test_extract_encrypted_pdf_is_rejected(tmp_path: Path) -> None:
    settings = build_settings(tmp_path)
    path = tmp_path / "encrypted.pdf"
    create_pdf(path, encrypted=True)
    with pytest.raises(ApiError) as exc:
        extract_document(path, "pdf", settings)
    assert exc.value.code == "PDF_ENCRYPTED"


def test_validate_docx_archive_rejects_missing_document_xml(tmp_path: Path) -> None:
    settings = build_settings(tmp_path)
    path = tmp_path / "bad.docx"
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr("[Content_Types].xml", "<Types />")
    with pytest.raises(ApiError):
        validate_docx_archive(path, settings)


def test_validate_docx_archive_rejects_entry_limit(tmp_path: Path) -> None:
    settings = build_settings(tmp_path, MAX_DOCX_ZIP_ENTRIES=10)
    path = tmp_path / "limit.docx"
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr("[Content_Types].xml", "a")
        archive.writestr("word/document.xml", "b")
        for index in range(9):
            archive.writestr(f"extra-{index}.bin", "c")
    with pytest.raises(ApiError) as exc:
        validate_docx_archive(path, settings)
    assert exc.value.code == "UNSAFE_DOCX_ARCHIVE"


def test_validate_docx_archive_rejects_traversal(tmp_path: Path) -> None:
    settings = build_settings(tmp_path)
    path = tmp_path / "bad.docx"
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr("[Content_Types].xml", "a")
        archive.writestr("word/document.xml", "b")
        archive.writestr("../evil.txt", "c")
    with pytest.raises(ApiError) as exc:
        validate_docx_archive(path, settings)
    assert exc.value.code == "UNSAFE_DOCX_ARCHIVE"


def test_validate_docx_archive_rejects_corrupt_zip(tmp_path: Path) -> None:
    settings = build_settings(tmp_path)
    path = tmp_path / "bad.docx"
    path.write_bytes(b"not-a-zip")
    with pytest.raises(ApiError) as exc:
        validate_docx_archive(path, settings)
    assert exc.value.code == "INVALID_DOCX"


def test_extract_document_isolated_success(tmp_path: Path) -> None:
    settings = build_settings(tmp_path)
    path = tmp_path / "sample.txt"
    path.write_text("hello", encoding="utf-8")
    extracted = extract_document_isolated(path, "txt", settings, process_factory=InlineProcess, poll_interval_seconds=0.0)
    assert extracted.text == "hello"


def test_extract_document_isolated_timeout(monkeypatch, tmp_path: Path) -> None:
    settings = build_settings(tmp_path, DOCUMENT_EXTRACTION_TIMEOUT_SECONDS=5)
    path = tmp_path / "sample.txt"
    path.write_text("hello", encoding="utf-8")
    ticks = iter([0.0, 10.0, 10.1])
    monkeypatch.setattr("app.documents.isolated_extraction.time.monotonic", lambda: next(ticks))
    with pytest.raises(ApiError) as exc:
        extract_document_isolated(path, "txt", settings, process_factory=HangingProcess, poll_interval_seconds=0.0)
    assert exc.value.code == "DOCUMENT_EXTRACTION_TIMEOUT"


def test_extract_document_isolated_memory_limit(tmp_path: Path) -> None:
    settings = build_settings(tmp_path, DOCUMENT_EXTRACTION_MEMORY_MB=128)
    path = tmp_path / "sample.txt"
    path.write_text("hello", encoding="utf-8")
    with pytest.raises(ApiError) as exc:
        extract_document_isolated(
            path,
            "txt",
            settings,
            process_factory=HangingProcess,
            memory_reader=lambda pid: 129 * 1024 * 1024,
            poll_interval_seconds=0.0,
        )
    assert exc.value.code == "DOCUMENT_EXTRACTION_MEMORY_LIMIT"


def test_extract_document_isolated_worker_exception(tmp_path: Path) -> None:
    settings = build_settings(tmp_path)
    path = tmp_path / "sample.bin"
    path.write_bytes(b"\x00\x00")
    with pytest.raises(ApiError) as exc:
        extract_document_isolated(path, "txt", settings, process_factory=InlineProcess, poll_interval_seconds=0.0)
    assert exc.value.code == "INVALID_TEXT_ENCODING"
